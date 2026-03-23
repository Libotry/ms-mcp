"""文档解析器模块。

支持多种文档格式的解析：
- PDF: 使用 PyMuPDF4LLM 转为 Markdown
- Markdown: 直接读取
- Word: 使用 python-docx
- HTML: 使用 BeautifulSoup
"""

import re
from abc import ABC, abstractmethod
from pathlib import Path
from typing import Literal

DocType = Literal["pdf", "md", "docx", "html", "txt"]


class DocumentParser(ABC):
    """文档解析器基类。"""

    @abstractmethod
    def parse(self, file_path: Path) -> str:
        """解析文档，返回 Markdown 格式内容。"""
        pass


class PDFParser(DocumentParser):
    """PDF 文档解析器，使用 PyMuPDF4LLM 转为 Markdown。"""

    def parse(self, file_path: Path) -> str:
        """解析 PDF 文件，返回 Markdown 内容。"""
        try:
            import pymupdf4llm
        except ImportError:
            raise ImportError(
                "请安装 pymupdf4llm: pip install pymupdf4llm"
            )

        if not file_path.exists():
            raise FileNotFoundError(f"文件不存在：{file_path}")

        # 转换为 Markdown
        md_text = pymupdf4llm.to_markdown(str(file_path))
        return md_text


class MarkdownParser(DocumentParser):
    """Markdown 文档解析器。"""

    def parse(self, file_path: Path) -> str:
        """读取 Markdown 文件内容。"""
        if not file_path.exists():
            raise FileNotFoundError(f"文件不存在：{file_path}")

        encodings = ["utf-8", "utf-8-sig", "gbk"]
        for enc in encodings:
            try:
                with open(file_path, "r", encoding=enc) as f:
                    return f.read()
            except (UnicodeDecodeError, UnicodeError):
                continue

        raise ValueError(f"无法解码文件：{file_path}")


class TextParser(DocumentParser):
    """纯文本文档解析器。"""

    def parse(self, file_path: Path) -> str:
        """读取文本文件内容。"""
        if not file_path.exists():
            raise FileNotFoundError(f"文件不存在：{file_path}")

        encodings = ["utf-8", "utf-8-sig", "gbk"]
        for enc in encodings:
            try:
                with open(file_path, "r", encoding=enc) as f:
                    return f.read()
            except (UnicodeDecodeError, UnicodeError):
                continue

        raise ValueError(f"无法解码文件：{file_path}")


class DocxParser(DocumentParser):
    """Word 文档解析器。"""

    def parse(self, file_path: Path) -> str:
        """解析 Word 文件，返回 Markdown 格式。"""
        try:
            from docx import Document
        except ImportError:
            raise ImportError("请安装 python-docx: pip install python-docx")

        if not file_path.exists():
            raise FileNotFoundError(f"文件不存在：{file_path}")

        doc = Document(str(file_path))
        paragraphs = []

        for para in doc.paragraphs:
            if para.text.strip():
                # 简单的样式判断
                if para.style.name.startswith("Heading"):
                    level = para.style.name.replace("Heading", "").strip()
                    try:
                        level_num = int(level)
                        prefix = "#" * level_num
                        paragraphs.append(f"{prefix} {para.text}")
                    except ValueError:
                        paragraphs.append(f"# {para.text}")
                else:
                    paragraphs.append(para.text)

        # 处理表格
        for table in doc.tables:
            md_table = []
            for i, row in enumerate(table.rows):
                cells = [cell.text.strip() for cell in row.cells]
                md_table.append("| " + " | ".join(cells) + " |")
                if i == 0:
                    separator = "| " + " | ".join(["---"] * len(cells)) + " |"
                    md_table.insert(1, separator)
            paragraphs.append("\n".join(md_table))

        return "\n\n".join(paragraphs)


class HTMLParser(DocumentParser):
    """HTML 文档解析器。"""

    def parse(self, file_path: Path) -> str:
        """解析 HTML 文件，返回 Markdown 格式。"""
        try:
            from bs4 import BeautifulSoup
        except ImportError:
            raise ImportError("请安装 beautifulsoup4: pip install beautifulsoup4")

        if not file_path.exists():
            raise FileNotFoundError(f"文件不存在：{file_path}")

        encodings = ["utf-8", "utf-8-sig", "gbk"]
        html_content = None
        for enc in encodings:
            try:
                with open(file_path, "r", encoding=enc) as f:
                    html_content = f.read()
                break
            except (UnicodeDecodeError, UnicodeError):
                continue

        if html_content is None:
            raise ValueError(f"无法解码文件：{file_path}")

        soup = BeautifulSoup(html_content, "html.parser")

        # 提取标题
        headings = []
        for tag in ["h1", "h2", "h3", "h4", "h5", "h6"]:
            for elem in soup.find_all(tag):
                level = int(tag[1])
                prefix = "#" * level
                headings.append(f"{prefix} {elem.get_text(strip=True)}")

        # 提取段落
        paragraphs = []
        for p in soup.find_all(["p", "li"]):
            text = p.get_text(strip=True)
            if text:
                paragraphs.append(text)

        # 提取表格
        tables = []
        for table in soup.find_all("table"):
            md_table = []
            rows = table.find_all("tr")
            for i, row in enumerate(rows):
                cells = []
                for cell in row.find_all(["td", "th"]):
                    cells.append(cell.get_text(strip=True))
                if cells:
                    md_table.append("| " + " | ".join(cells) + " |")
                    if i == 0:
                        separator = "| " + " | ".join(["---"] * len(cells)) + " |"
                        md_table.insert(1, separator)
            if md_table:
                tables.append("\n".join(md_table))

        # 组合所有内容
        all_parts = headings + paragraphs + tables
        return "\n\n".join(all_parts)


# ── 工厂函数 ──────────────────────────────────────────────


def get_parser(doc_type: DocType) -> DocumentParser:
    """根据文档类型获取对应的解析器。"""
    parsers = {
        "pdf": PDFParser,
        "md": MarkdownParser,
        "docx": DocxParser,
        "html": HTMLParser,
        "txt": TextParser,
    }

    if doc_type not in parsers:
        raise ValueError(f"不支持的文档类型：{doc_type}")

    return parsers[doc_type]()


def detect_doc_type(file_path: Path) -> DocType:
    """根据文件扩展名检测文档类型。"""
    suffix_map = {
        ".pdf": "pdf",
        ".md": "md",
        ".markdown": "md",
        ".docx": "docx",
        ".html": "html",
        ".htm": "html",
        ".txt": "txt",
    }

    suffix = file_path.suffix.lower()
    if suffix not in suffix_map:
        raise ValueError(f"不支持的文件格式：{suffix}")

    return suffix_map[suffix]  # type: ignore


def parse_document(file_path: Path) -> tuple[str, DocType]:
    """解析文档，返回 (Markdown 内容，文档类型)。"""
    if isinstance(file_path, str):
        file_path = Path(file_path)

    doc_type = detect_doc_type(file_path)
    parser = get_parser(doc_type)
    content = parser.parse(file_path)

    return content, doc_type


# ── 辅助函数 ──────────────────────────────────────────────


def extract_metadata(content: str, source_file: str) -> dict:
    """从 Markdown 内容中提取元数据（章节、页码等）。"""
    metadata = {
        "source": source_file,
        "chapters": [],
        "sections": [],
    }

    # 提取章节标题
    heading_pattern = re.compile(r"^(#{1,6})\s+(.+)$", re.MULTILINE)
    for match in heading_pattern.finditer(content):
        level = len(match.group(1))
        title = match.group(2).strip()

        if level == 1:
            metadata["chapters"].append(title)
        elif level <= 3:
            metadata["sections"].append({"level": level, "title": title})

    return metadata
